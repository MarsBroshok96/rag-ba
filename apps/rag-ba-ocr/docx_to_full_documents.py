from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document

def extract_blocks_from_docx(doc_path: Path) -> list[dict[str, Any]]:
    doc = Document(doc_path)

    blocks: list[dict[str, Any]] = []
    region_counter = 1

    for element in doc.element.body:
        # Paragraph
        if element.tag.endswith("p"):
            p = element
            para = next((x for x in doc.paragraphs if x._p == p), None)
            if not para:
                continue

            text = para.text.strip()
            if not text:
                continue

            style = para.style.name.lower() if para.style else ""

            if "heading" in style:
                typ = "title"
            else:
                typ = "text"

            blocks.append(
                {
                    "region_id": f"{doc_path.stem}__r{region_counter}",
                    "type": typ,
                    "text": text,
                }
            )
            region_counter += 1

        # Table
        elif element.tag.endswith("tbl"):
            tbl = next((x for x in doc.tables if x._tbl == element), None)
            if not tbl:
                continue

            rows_text = []
            for row in tbl.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(row_cells))

            table_text = "\n".join(rows_text).strip()
            if table_text:
                blocks.append(
                    {
                        "region_id": f"{doc_path.stem}__r{region_counter}",
                        "type": "table",
                        "text": table_text,
                    }
                )
                region_counter += 1

    return blocks


def main() -> None:
    # rag-ba root (соседний репозиторий)
    rag_ba_root = (Path(__file__).resolve().parent / ".." / "..").resolve()

    # unified inbox
    input_dir = rag_ba_root / "data" / "inbox" / "word"
    assert input_dir.exists(), f"Word inbox not found: {input_dir}"
    
    output_root = Path("export")
    output_root.mkdir(parents=True, exist_ok=True)

    docx_files = sorted(input_dir.glob("*.docx"))
    assert docx_files, "No .docx files found in data/inbox/docx"

    for doc_path in docx_files:
        print(f"Processing {doc_path.name}")

        blocks = extract_blocks_from_docx(doc_path)

        out_data = {
            "doc_id": doc_path.stem,
            "source_file": doc_path.name,
            "source_path": str(doc_path.resolve()),
            "page": 1,  # docx считаем одной "страницей"
            "blocks": blocks,
        }
        doc_out_dir = output_root / doc_path.stem
        doc_out_dir.mkdir(parents=True, exist_ok=True)
        
        out_path = doc_out_dir / "full_document.json"
        out_path.write_text(
            json.dumps(out_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        print(f"Saved → {out_path}")

    print("Done.")
    

if __name__ == "__main__":
    main()